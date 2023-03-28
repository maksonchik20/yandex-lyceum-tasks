from docx import Document
import sys

document = Document()
data = [line.rstrip('\n') for line in sys.stdin]
place = data[0]
date = data[1]
for i in range(2, len(data)):
    document.add_heading(f'Добрый день, ' + data[i] + ', 0')
    document.add_paragraph(f'Приглашаю тебя на праздник 8 мара, которое состоится { place } { date }!')
    document.add_page_break()
document.save('test.docx')